#!/usr/bin/env python3
"""
Exhibit Index Generator — CLI tool for litigators
Scans a directory of exhibit PDFs, extracts metadata (page count, Bates from filenames/OCR),
outputs court-ready index as Markdown, CSV, or DOCX.
"""
import os, re, csv, json, sys, argparse, subprocess
from pathlib import Path
from datetime import datetime

try:
    import fitz  # pymupdf
except ImportError:
    print("Error: pymupdf not installed. Run: pip install pymupdf", file=sys.stderr)
    sys.exit(1)

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
except ImportError:
    Document = None

# ---------- Bates extraction ----------
BATES_PATTERNS = [
    re.compile(r'(?i)(?:bates?|exh?\.?)\s*[:#]?\s*([A-Z]{2,6}[-_]?\d{4,8})'),
    re.compile(r'(?i)([A-Z]{2,6}[-_]\d{4,8})'),  # e.g. EXH-00001, DOC_001234
    re.compile(r'(?i)(?:^|[_\-\s])([A-Z]{2,5}\d{4,8})(?:[_\-\s]|$)'),  # prefix + digits
]

def extract_bates_from_filename(fname):
    for pat in BATES_PATTERNS:
        m = pat.search(fname)
        if m:
            return m.group(1).upper().replace('_', '-')
    return None

def extract_bates_from_pdf_first_page(pdf_path):
    """OCR-free: try text layer on page 1 for Bates stamp"""
    try:
        doc = fitz.open(pdf_path)
        if len(doc) > 0:
            text = doc[0].get_text()
            for pat in BATES_PATTERNS:
                m = pat.search(text)
                if m:
                    doc.close()
                    return m.group(1).upper().replace('_', '-')
        doc.close()
    except Exception:
        pass
    return None

# ---------- PDF metadata ----------
def get_pdf_metadata(pdf_path):
    """Return (page_count, file_size_bytes, bates_from_text)"""
    try:
        doc = fitz.open(pdf_path)
        pages = len(doc)
        doc.close()
        size = os.path.getsize(pdf_path)
        bates = extract_bates_from_pdf_first_page(pdf_path)
        return pages, size, bates
    except Exception as e:
        return 0, os.path.getsize(pdf_path) if os.path.exists(pdf_path) else 0, None

# ---------- Natural sort for exhibit numbers ----------
def exhibit_sort_key(path):
    """Sort by exhibit number in filename: Exh 1, Exh 2, Exh 10..."""
    fname = path.name
    # Try to find exhibit number
    m = re.search(r'(?i)(?:exh?\.?|exhibit)\s*[:#]?\s*(\d+)', fname)
    if m:
        return (0, int(m.group(1)), fname)
    m = re.search(r'(?i)^(\d+)', fname)
    if m:
        return (0, int(m.group(1)), fname)
    return (1, fname.lower(), fname)

# ---------- Output formatters ----------
def write_markdown(exhibits, out_path, case_name="", case_no="", court=""):
    with open(out_path, 'w') as f:
        f.write(f"# Exhibit Index\n\n")
        if case_name: f.write(f"**Case:** {case_name}\n")
        if case_no: f.write(f"**Case No.:** {case_no}\n")
        if court: f.write(f"**Court:** {court}\n")
        f.write(f"**Generated:** {datetime.now().strftime('%Y-%m-%d %H:%M')}\n")
        f.write(f"**Total Exhibits:** {len(exhibits)}\n")
        f.write(f"**Total Pages:** {sum(e['pages'] for e in exhibits)}\n\n")
        f.write("| Ex. # | Bates / ID | Description | Pages | Size |\n")
        f.write("|-------|------------|-------------|-------|------|\n")
        for i, e in enumerate(exhibits, 1):
            desc = e['description'].replace('|', '\\|')
            bates = e['bates'] or '—'
            f.write(f"| {i} | {bates} | {desc} | {e['pages']} | {e['size_kb']} KB |\n")
    print(f"✅ Markdown written to {out_path}")

def write_csv(exhibits, out_path, case_name="", case_no="", court=""):
    with open(out_path, 'w', newline='') as f:
        w = csv.writer(f)
        w.writerow(["Exhibit #", "Bates / ID", "Description", "Pages", "Size (KB)", "File Path"])
        for i, e in enumerate(exhibits, 1):
            w.writerow([i, e['bates'] or '', e['description'], e['pages'], e['size_kb'], e['path']])
    print(f"✅ CSV written to {out_path}")

def write_docx(exhibits, out_path, case_name="", case_no="", court=""):
    if Document is None:
        print("Error: python-docx not installed. Run: pip install python-docx", file=sys.stderr)
        return False
    doc = Document()
    # Styles
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    # Title
    title = doc.add_heading('Exhibit Index', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Meta table
    meta = doc.add_table(rows=4, cols=2)
    meta.style = 'Light Shading Accent 1'
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_rows = [
        ("Case", case_name or "—"),
        ("Case No.", case_no or "—"),
        ("Court", court or "—"),
        ("Generated", datetime.now().strftime('%Y-%m-%d %H:%M')),
    ]
    for i, (k, v) in enumerate(meta_rows):
        meta.rows[i].cells[0].text = k
        meta.rows[i].cells[1].text = v
    doc.add_paragraph()
    # Summary
    doc.add_paragraph(f"Total Exhibits: {len(exhibits)}")
    doc.add_paragraph(f"Total Pages: {sum(e['pages'] for e in exhibits)}")
    doc.add_paragraph()
    # Exhibit table
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Light Shading Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    hdr[0].text = "Ex. #"
    hdr[1].text = "Bates / ID"
    hdr[2].text = "Description"
    hdr[3].text = "Pages"
    hdr[4].text = "Size (KB)"
    for e in exhibits:
        row = table.add_row().cells
        row[0].text = str(exhibits.index(e) + 1)
        row[1].text = e['bates'] or '—'
        row[2].text = e['description']
        row[3].text = str(e['pages'])
        row[4].text = str(e['size_kb'])
    # Set column widths
    for row in table.rows:
        row.cells[0].width = Cm(1.5)
        row.cells[1].width = Cm(3.5)
        row.cells[2].width = Cm(10)
        row.cells[3].width = Cm(1.5)
        row.cells[4].width = Cm(2)
    doc.save(out_path)
    print(f"✅ DOCX written to {out_path}")
    return True

# ---------- Main ----------
def main():
    ap = argparse.ArgumentParser(
        description="Exhibit Index Generator — builds court-ready exhibit indexes from PDF folders",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  %(prog)s /path/to/exhibits -o index.md
  %(prog)s /path/to/exhibits -o index.csv --case "State v. Smith" --case-no "CR-2026-00123"
  %(prog)s /path/to/exhibits -o index.docx --court "High Court of Delhi"
  %(prog)s /path/to/exhibits --recursive -o index.md
        """
    )
    ap.add_argument("input_dir", help="Directory containing exhibit PDFs")
    ap.add_argument("-o", "--output", required=True, help="Output file (.md, .csv, or .docx)")
    ap.add_argument("-r", "--recursive", action="store_true", help="Scan subdirectories recursively")
    ap.add_argument("--case", default="", help="Case name (e.g., 'State v. Smith')")
    ap.add_argument("--case-no", default="", dest="case_no", help="Case number")
    ap.add_argument("--court", default="", help="Court name")
    ap.add_argument("--no-bates", action="store_true", help="Skip Bates extraction (faster)")
    args = ap.parse_args()

    in_dir = Path(args.input_dir)
    if not in_dir.is_dir():
        print(f"Error: {in_dir} is not a directory", file=sys.stderr)
        sys.exit(1)

    # Find PDFs
    pattern = "**/*.pdf" if args.recursive else "*.pdf"
    pdf_files = sorted(in_dir.glob(pattern), key=exhibit_sort_key)
    if not pdf_files:
        print("No PDF files found.", file=sys.stderr)
        sys.exit(1)

    print(f"Scanning {len(pdf_files)} PDF(s)...")

    exhibits = []
    for pdf in pdf_files:
        pages, size, bates_pdf = get_pdf_metadata(pdf)
        bates_file = extract_bates_from_filename(pdf.name)
        bates = bates_file or bates_pdf
        if args.no_bates:
            bates = None

        # Description: filename without extension, cleaned
        desc = pdf.stem
        desc = re.sub(r'(?i)^(exh?\.?|exhibit)\s*[:#]?\s*\d+\s*[-_]?\s*', '', desc)
        desc = desc.replace('_', ' ').replace('-', ' ').strip()
        desc = desc[:120]  # reasonable cap

        exhibits.append({
            "path": str(pdf),
            "description": desc,
            "pages": pages,
            "size_kb": round(size / 1024, 1),
            "bates": bates,
        })

    # Output
    ext = Path(args.output).suffix.lower()
    if ext in ('.md', '.markdown'):
        write_markdown(exhibits, args.output, args.case, args.case_no, args.court)
    elif ext == '.csv':
        write_csv(exhibits, args.output, args.case, args.case_no, args.court)
    elif ext == '.docx':
        ok = write_docx(exhibits, args.output, args.case, args.case_no, args.court)
        if not ok:
            sys.exit(1)
    else:
        print(f"Error: unsupported output format '{ext}'. Use .md, .csv, or .docx", file=sys.stderr)
        sys.exit(1)

    print(f"Done. {len(exhibits)} exhibits indexed.")

if __name__ == "__main__":
    main()
